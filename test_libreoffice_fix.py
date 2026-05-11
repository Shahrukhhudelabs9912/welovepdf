#!/usr/bin/env python3
"""
Test LibreOffice Word to PDF conversion.
"""

import subprocess
import tempfile
import os
from pathlib import Path
import sys

def create_test_docx():
    """Create a simple test DOCX file using Python."""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Test Document for PDF Conversion', 0)
        doc.add_paragraph('This is a test document to verify Word to PDF conversion.')
        doc.add_paragraph('Created by WeLovePDF Word to PDF converter.')
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_path = f.name
            doc.save(temp_path)
            return temp_path
    except ImportError:
        # If python-docx is not installed, create a simple text file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False, mode='w') as f:
            temp_path = f.name
            # Just create an empty file for testing
            return temp_path

def test_libreoffice_conversion():
    """Test if LibreOffice can convert Word to PDF."""
    print("Testing LibreOffice Word to PDF conversion...")
    
    # Create test document
    test_docx = create_test_docx()
    print(f"Created test document: {test_docx}")
    
    # Create temp directory for output
    with tempfile.TemporaryDirectory() as temp_dir:
        print(f"Temp output directory: {temp_dir}")
        
        # Try to convert using LibreOffice
        libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
        
        if not os.path.exists(libreoffice_path):
            print(f"ERROR: LibreOffice not found at {libreoffice_path}")
            return False
        
        try:
            # Run LibreOffice conversion
            cmd = [
                libreoffice_path,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', temp_dir,
                test_docx
            ]
            
            print(f"Running command: {' '.join(cmd)}")
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=30
            )
            
            print(f"Return code: {result.returncode}")
            print(f"Stdout: {result.stdout}")
            print(f"Stderr: {result.stderr}")
            
            if result.returncode == 0:
                # Check for converted PDF
                pdf_files = list(Path(temp_dir).glob('*.pdf'))
                if pdf_files:
                    pdf_file = pdf_files[0]
                    print(f"SUCCESS: Converted PDF created: {pdf_file}")
                    print(f"PDF file size: {os.path.getsize(pdf_file)} bytes")
                    return True
                else:
                    print("ERROR: No PDF file was created")
                    return False
            else:
                print("ERROR: LibreOffice conversion failed")
                return False
                
        except subprocess.TimeoutExpired:
            print("ERROR: LibreOffice conversion timed out")
            return False
        except Exception as e:
            print(f"ERROR: {e}")
            return False
        finally:
            # Clean up test document
            if os.path.exists(test_docx):
                os.unlink(test_docx)
                print(f"Cleaned up test document: {test_docx}")

def test_python_docx_import():
    """Test if python-docx is available."""
    try:
        from docx import Document
        print("SUCCESS: python-docx is installed")
        return True
    except ImportError:
        print("WARNING: python-docx is not installed")
        print("Install with: pip install python-docx")
        return False

if __name__ == "__main__":
    print("=" * 60)
    print("LibreOffice Conversion Test")
    print("=" * 60)
    
    # Test python-docx
    test_python_docx_import()
    
    print("\n" + "-" * 60)
    
    # Test LibreOffice conversion
    success = test_libreoffice_conversion()
    
    print("\n" + "=" * 60)
    if success:
        print("SUCCESS: LibreOffice Word to PDF conversion works!")
    else:
        print("FAILURE: LibreOffice conversion test failed")
    print("=" * 60)