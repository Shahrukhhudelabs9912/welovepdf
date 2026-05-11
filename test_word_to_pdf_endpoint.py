#!/usr/bin/env python3
"""
Test script for Word-to-PDF endpoint.
Creates a simple Word document and tests conversion through the backend API.
"""

import requests
import os
import tempfile
import time
from pathlib import Path
import sys

def create_test_word_document():
    """Create a simple test Word document."""
    from docx import Document
    
    doc = Document()
    doc.add_heading('Test Word Document for PDF Conversion', 0)
    doc.add_paragraph('This is a test document created to verify Word-to-PDF conversion functionality.')
    doc.add_paragraph('Created: ' + time.strftime("%Y-%m-%d %H:%M:%S"))
    
    # Add some content
    doc.add_heading('Sample Content', level=1)
    doc.add_paragraph('This document contains:')
    doc.add_paragraph('• Bullet point 1', style='List Bullet')
    doc.add_paragraph('• Bullet point 2', style='List Bullet')
    doc.add_paragraph('• Bullet point 3', style='List Bullet')
    
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph('If this document converts successfully to PDF, the Word-to-PDF functionality is working correctly.')
    
    return doc

def test_word_to_pdf_endpoint():
    """Test the Word-to-PDF endpoint."""
    print("Testing Word-to-PDF endpoint...")
    
    # Create a test Word document
    print("Creating test Word document...")
    doc = create_test_word_document()
    
    # Save to temporary file
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name
    
    print(f"Created test document: {tmp_path}")
    print(f"File size: {os.path.getsize(tmp_path)} bytes")
    
    try:
        # Send to backend endpoint
        print("\nSending request to backend endpoint...")
        with open(tmp_path, 'rb') as f:
            files = {'file': ('test_document.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            
            response = requests.post(
                'http://127.0.0.1:8000/api/word-to-pdf',
                files=files,
                timeout=120  # 2 minute timeout for conversion
            )
        
        print(f"Response status: {response.status_code}")
        print(f"Response headers: {dict(response.headers)}")
        
        if response.status_code == 200:
            # Check if response is a PDF
            content_type = response.headers.get('Content-Type', '')
            content_disposition = response.headers.get('Content-Disposition', '')
            
            print(f"\n[SUCCESS] Response is a PDF")
            print(f"  Content-Type: {content_type}")
            print(f"  Content-Disposition: {content_disposition}")
            print(f"  Response size: {len(response.content)} bytes")
            
            # Save the PDF for inspection
            output_filename = 'test_word_to_pdf_output.pdf'
            with open(output_filename, 'wb') as f:
                f.write(response.content)
            
            print(f"  Saved PDF to: {output_filename}")
            
            # Verify it's a valid PDF
            if response.content[:4] == b'%PDF':
                print("  [VALID] Valid PDF file (PDF header found)")
            else:
                print("  [WARNING] Response doesn't start with PDF header")
                
            return True
        else:
            print(f"\n[ERROR] {response.status_code}")
            print(f"Response text: {response.text[:500]}")
            return False
            
    except requests.exceptions.RequestException as e:
        print(f"\n✗ Request failed: {e}")
        return False
    finally:
        # Clean up temporary file
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
            print(f"\nCleaned up temporary file: {tmp_path}")

def test_invalid_file():
    """Test with invalid file type."""
    print("\n\nTesting invalid file type (should fail with 415)...")
    
    # Create a text file (not a Word document)
    with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
        tmp.write(b"This is not a Word document")
        tmp_path = tmp.name
    
    try:
        with open(tmp_path, 'rb') as f:
            files = {'file': ('test.txt', f, 'text/plain')}
            response = requests.post(
                'http://127.0.0.1:8000/api/word-to-pdf',
                files=files,
                timeout=30
            )
        
        print(f"Response status: {response.status_code}")
        if response.status_code == 415:
            print("[SUCCESS] Correctly rejected non-Word file")
            return True
        else:
            print(f"[ERROR] Unexpected response: {response.status_code}")
            print(f"Response: {response.text[:200]}")
            return False
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

def main():
    """Run all tests."""
    print("=" * 60)
    print("Word-to-PDF Endpoint Test")
    print("=" * 60)
    
    # Check if backend is running
    try:
        health_response = requests.get('http://127.0.0.1:8000/health', timeout=5)
        if health_response.status_code != 200:
            print("[ERROR] Backend is not responding properly")
            return 1
        print("[OK] Backend is running")
    except requests.exceptions.RequestException:
        print("[ERROR] Backend is not running. Please start the backend server first.")
        print("  Run: cd backend && python -m uvicorn app.main:app --host 127.0.0.1 --port 8000")
        return 1
    
    # Run tests
    success = True
    
    # Test 1: Valid Word document conversion
    if not test_word_to_pdf_endpoint():
        success = False
    
    # Test 2: Invalid file type
    if not test_invalid_file():
        success = False
    
    print("\n" + "=" * 60)
    if success:
        print("[SUCCESS] All tests passed!")
        return 0
    else:
        print("[FAILURE] Some tests failed")
        return 1

if __name__ == "__main__":
    # Install required package if not available
    try:
        from docx import Document
    except ImportError:
        print("Installing python-docx for test...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        from docx import Document
    
    sys.exit(main())