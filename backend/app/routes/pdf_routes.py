"""
PDF processing routes for PDFOrca API.
"""
import asyncio
import io
import os
import logging
import traceback
from pathlib import Path
from typing import List, Optional
from fastapi import APIRouter, UploadFile, File, Form, HTTPException, status
from fastapi.responses import StreamingResponse, Response

from app.services.pdf_service import PDFService, ImageToPDFService, PDFToImageService
from app.services.organize_pdf_service import OrganizePDFService
from app.utils import (
    validate_file_type,
    validate_file_size,
    read_upload_file,
    create_file_response,
    create_zip_response,
    validate_uploaded_files,
    handle_pdf_error,
    PDFProcessingError,
    run_blocking,
    heavy_job_slot,
    resolve_libreoffice_path,
)
from app.config import settings
from app.schemas import (
    PlaceholderResponse,
)

logger = logging.getLogger(__name__)

router = APIRouter()


@router.post("/merge-pdf", summary="Merge multiple PDF files")
async def merge_pdf(files: List[UploadFile] = File(...)):
    """
    Merge multiple PDF files into a single PDF.
    Handles AES-encrypted PDFs gracefully; password-protected PDFs
    return a clear user-facing error.
    
    - **files**: List of PDF files to merge (2-10 files)
    
    Returns merged PDF file for download.
    """
    logger.info(f"Merge PDF request: {len(files)} file(s)")
    
    try:
        # Validate files
        validate_uploaded_files(files, max_files=10)
        
        pdf_files = []
        for file in files:
            # Validate file type
            if not validate_file_type(file, settings.ALLOWED_PDF_TYPES):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"File {file.filename} is not a valid PDF. Allowed types: {settings.ALLOWED_PDF_TYPES}"
                )
            
            # Validate file size
            if not validate_file_size(file):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"File {file.filename} exceeds maximum size of {settings.MAX_UPLOAD_SIZE} bytes"
                )
            
            # Read file
            file_bytes = await read_upload_file(file)
            pdf_files.append(file_bytes)
            logger.debug(f"Read file: {file.filename} ({len(file_bytes)} bytes)")
        
        # Merge PDFs
        logger.info(f"Merging {len(pdf_files)} PDF(s)...")
        merged_pdf = await run_blocking(PDFService.merge_pdfs, pdf_files)
        logger.info(f"Merge successful: {len(merged_pdf)} bytes output")
        
        # Return merged PDF
        return create_file_response(
            merged_pdf,
            filename="merged.pdf",
            media_type="application/pdf"
        )
        
    except HTTPException:
        raise
    except PDFProcessingError as e:
        logger.warning(f"Merge PDF processing error: {e.message}")
        return handle_pdf_error(e)
    except Exception as e:
        logger.error(f"Unexpected merge PDF error: {str(e)}", exc_info=True)
        return handle_pdf_error(e)


@router.post("/split-pdf", summary="Split PDF file")
async def split_pdf(
    file: UploadFile = File(...),
    split_method: str = Form("all", description="Split method: all, range, every, or pages"),
    page_range: Optional[str] = Form(None, description="Page ranges e.g. '1-5,8-10' (range method)"),
    pages_per_split: Optional[int] = Form(None, description="Pages per split file (every method)"),
    specific_pages: Optional[str] = Form(None, description="Specific pages e.g. '1,3,5-7' (pages method)"),
    output_format: str = Form("individual", description="Output format: individual or single"),
    naming_pattern: str = Form("page_{n}.pdf", description="Naming pattern with {n} placeholder"),
):
    """
    Split a PDF file using various methods.

    - **split_method**: "all" (every page separate), "range" (page ranges), "every" (every N pages), "pages" (specific pages)
    - **page_range**: Comma-separated ranges e.g. "1-5,8-10"
    - **pages_per_split**: Number of pages per output file
    - **specific_pages**: Comma-separated pages and ranges e.g. "1,3,5-7"
    - **output_format**: "individual" (separate files as ZIP) or "single" (single PDF)
    - **naming_pattern**: Filename template, use {n} for part number
    """
    logger.info(
        f"Split PDF request: file={file.filename}, method={split_method}, "
        f"range={page_range}, every={pages_per_split}, pages={specific_pages}, "
        f"format={output_format}, naming={naming_pattern}"
    )
    try:
        # Validate file
        if not validate_file_type(file, settings.ALLOWED_PDF_TYPES):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File {file.filename} is not a valid PDF. Allowed types: {settings.ALLOWED_PDF_TYPES}"
            )

        if not validate_file_size(file):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File {file.filename} exceeds maximum size of {settings.MAX_UPLOAD_SIZE} bytes"
            )

        # Read file
        pdf_bytes = await read_upload_file(file)

        # Split PDF
        pages = await run_blocking(
            PDFService.split_pdf,
            pdf_bytes,
            split_method=split_method,
            page_range=page_range,
            pages_per_split=pages_per_split,
            specific_pages=specific_pages,
            output_format=output_format,
            naming_pattern=naming_pattern,
        )

        if not pages:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No pages produced from split operation"
            )

        logger.info(f"Split PDF produced {len(pages)} output file(s)")

        # Return single PDF or ZIP based on output format
        if output_format == "single" and len(pages) == 1:
            filename, file_bytes = pages[0]
            return create_file_response(
                file_bytes,
                filename=filename,
            )
        else:
            base_name = file.filename.replace(".pdf", "") if file.filename else "split"
            return create_zip_response(
                pages,
                zip_filename=f"{base_name}_split.zip",
            )

    except HTTPException:
        raise
    except PDFProcessingError as e:
        logger.error(f"Split PDF processing error: {str(e)}")
        return handle_pdf_error(e)
    except Exception as e:
        logger.error(f"Split PDF error: {str(e)}")
        return handle_pdf_error(e)


@router.post("/jpg-to-pdf", summary="Convert images to PDF")
async def jpg_to_pdf(files: List[UploadFile] = File(...)):
    """
    Convert multiple JPG/PNG images to a single PDF.
    
    - **files**: List of image files to convert
    
    Returns PDF file containing all images.
    """
    try:
        # Validate files
        validate_uploaded_files(files, max_files=20)
        
        image_files = []
        for file in files:
            # Validate file type
            if not validate_file_type(file, settings.ALLOWED_IMAGE_TYPES):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"File {file.filename} is not a valid image. Allowed types: {settings.ALLOWED_IMAGE_TYPES}"
                )
            
            # Validate file size
            if not validate_file_size(file):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"File {file.filename} exceeds maximum size of {settings.MAX_UPLOAD_SIZE} bytes"
                )
            
            # Read file
            file_bytes = await read_upload_file(file)
            image_files.append(file_bytes)
        
        # Convert images to PDF
        pdf_bytes = await run_blocking(ImageToPDFService.convert_images_to_pdf, image_files)
        
        # Return PDF
        return create_file_response(
            pdf_bytes,
            filename="converted.pdf",
            media_type="application/pdf"
        )
        
    except HTTPException as http_exc:
        # Log the HTTP exception for debugging
        logger.debug(f"[DEBUG] HTTPException in jpg_to_pdf: {http_exc.status_code} - {http_exc.detail}")
        raise
    except PDFProcessingError as e:
        logger.error(f"JPG to PDF processing error: {str(e)}")
        return handle_pdf_error(e)
    except Exception as e:
        # Log the full error for debugging
        logger.error(f"JPG to PDF error: {type(e).__name__}: {str(e)}")
        return handle_pdf_error(e)


@router.post("/pdf-to-jpg", summary="Convert PDF to image(s)")
async def pdf_to_jpg(
    file: UploadFile = File(...),
    page_number: int = Form(0, ge=0, description="Page number to convert (0 = all pages, 1+ = specific page)"),
    quality: int = Form(85, ge=1, le=100, description="Image quality (1-100%)"),
    dpi: int = Form(150, ge=72, le=300, description="Image resolution in DPI")
):
    """
    Convert PDF pages to JPEG image(s).
    
    - **file**: PDF file to convert
    - **page_number**: 0 = convert all pages (returns ZIP), 1+ = convert specific page (returns single JPG)
    - **quality**: JPEG quality 1-100 (default 85)
    - **dpi**: Resolution in DPI (default 150)
    
    Returns single JPEG or ZIP containing all page images.
    """
    try:
        # Log incoming request details for debugging
        logger.debug(f"[DEBUG] pdf_to_jpg called: filename={file.filename}, page_number={page_number}, quality={quality}, dpi={dpi}")
        
        # Validate file
        if not validate_file_type(file, settings.ALLOWED_PDF_TYPES):
            error_msg = f"File {file.filename} is not a valid PDF. Allowed types: {settings.ALLOWED_PDF_TYPES}, got: {file.content_type}"
            logger.debug(f"[DEBUG] Validation failed: {error_msg}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=error_msg
            )
        
        if not validate_file_size(file):
            error_msg = f"File {file.filename} exceeds maximum size of {settings.MAX_UPLOAD_SIZE} bytes"
            logger.debug(f"[DEBUG] Size validation failed: {error_msg}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=error_msg
            )
        
        # Read file
        pdf_bytes = await read_upload_file(file)
        logger.debug(f"[DEBUG] Read PDF file: {len(pdf_bytes)} bytes")

        # Safety: predict pixel count and cap DPI if needed. Prevents the
        # Pillow decompression-bomb error on huge/high-DPI renders and keeps
        # render time bounded.
        effective_dpi, dpi_was_adjusted = PDFToImageService.safe_dpi(pdf_bytes, dpi)
        if dpi_was_adjusted:
            logger.debug(f"[DEBUG] DPI auto-capped {dpi} -> {effective_dpi} for safety")
        # Header passed through to the client so the UI can show a notice.
        # Single source of truth — built once and merged into every response.
        adjust_headers = {
            "X-DPI-Requested": str(dpi),
            "X-DPI-Used": str(effective_dpi),
            "X-DPI-Adjusted": "true" if dpi_was_adjusted else "false",
        }

        # Clean base filename (remove .pdf extension)
        base_name = file.filename.replace(".pdf", "").replace(".PDF", "")

        # Conversion is CPU-bound and synchronous (Poppler subprocess + Pillow
        # encode). Run it on the thread pool so the FastAPI event loop stays
        # responsive for other users while this single request is processing.
        loop = asyncio.get_running_loop()

        if page_number == 0:
            # Convert ALL pages → ZIP
            logger.debug(f"[DEBUG] All pages mode: converting all pages to ZIP (dpi={effective_dpi})")
            try:
                zip_bytes, zip_filename = await loop.run_in_executor(
                    None,
                    PDFToImageService.convert_pages_to_zip,
                    pdf_bytes, base_name, quality, effective_dpi,
                )
                logger.debug(f"[DEBUG] ZIP created: {zip_filename} ({len(zip_bytes)} bytes)")
                # Plain Response (not StreamingResponse). Starlette's
                # BaseHTTPMiddleware buffers streaming bodies chunk-by-chunk
                # and adds ~600ms per MB on this stack — for 10MB outputs that
                # was 6+ seconds of pure overhead. The bytes are already in
                # memory so a single-send Response is strictly faster.
                return Response(
                    content=zip_bytes,
                    media_type="application/zip",
                    headers={
                        "Content-Disposition": f'attachment; filename="{zip_filename}"',
                        "Content-Length": str(len(zip_bytes)),
                        **adjust_headers,
                    }
                )
            except Exception as zip_err:
                logger.debug(f"[DEBUG] ZIP conversion failed: {zip_err}, falling back to single page")
                # Fallback: convert only first page
                image_bytes = await loop.run_in_executor(
                    None,
                    PDFToImageService.convert_pdf_to_image,
                    pdf_bytes, 0, quality, effective_dpi,
                )
                return Response(
                    content=image_bytes,
                    media_type="image/jpeg",
                    headers={
                        "Content-Disposition": f'attachment; filename="{base_name}_page_1.jpg"',
                        "Content-Length": str(len(image_bytes)),
                        **adjust_headers,
                    }
                )
        else:
            # Convert single page
            page_idx = page_number - 1  # Convert from 1-indexed (user-facing) to 0-indexed
            logger.debug(f"[DEBUG] Single page mode: converting page {page_number} (index {page_idx})")

            # Validate page number
            total_pages = PDFToImageService._get_page_count(pdf_bytes)
            if total_pages > 0 and page_idx >= total_pages:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Page number {page_number} out of range. PDF has {total_pages} pages."
                )

            image_bytes = await loop.run_in_executor(
                None,
                PDFToImageService.convert_pdf_to_image,
                pdf_bytes, page_idx, quality, effective_dpi,
            )
            logger.debug(f"[DEBUG] Conversion successful: {len(image_bytes)} bytes")

            return Response(
                content=image_bytes,
                media_type="image/jpeg",
                headers={
                    "Content-Disposition": f'attachment; filename="{base_name}_page_{page_number}.jpg"',
                    "Content-Length": str(len(image_bytes)),
                    **adjust_headers,
                }
            )
        
    except HTTPException:
        raise
    except Exception as e:
        # No silent placeholder JPEG — that masked real failures (Poppler
        # missing, OOM, corrupt PDF) behind a "successful" white image and made
        # support tickets impossible to trace. Surface the error so the UI can
        # show it and Sentry/logs capture it.
        logger.error(f"pdf_to_jpg failed for {file.filename}: {type(e).__name__}: {e}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"PDF to JPG conversion failed: {str(e)[:200]}",
        )


@router.post("/add-watermark", summary="Add watermark to PDF")
async def add_watermark(
    file: UploadFile = File(...),
    watermark_type: str = Form("text", description="Type of watermark: 'text' or 'image'"),
    watermark_text: str = Form("", description="Watermark text (required for text watermarks)"),
    position: str = Form("diagonal", description="Watermark position: 'diagonal', 'center', 'top-left', 'top-right', 'bottom-left', 'bottom-right'"),
    opacity: int = Form(30, description="Watermark opacity (0-100)"),
    rotation: int = Form(45, description="Watermark rotation angle (0-360)"),
    pages: str = Form("all", description="Pages to apply watermark: 'all', 'first', 'last', 'custom'"),
    custom_page_range: str = Form("", description="Custom page range (e.g., '1,3,5-7') when pages='custom'"),
    font_size: int = Form(36, description="Font size in points (default: 36)"),
    color: str = Form("#808080", description="Color in hex format (default: #808080 gray)"),
    watermark_image: Optional[UploadFile] = File(None, description="Image file for image watermark (PNG/JPG)")
):
    """
    Add text or image watermark to PDF pages with configurable settings.
    
    - **file**: PDF file to watermark
    - **watermark_type**: Type of watermark: 'text' or 'image'
    - **watermark_text**: Text for watermark (required for text watermarks)
    - **position**: Watermark position
    - **opacity**: Watermark opacity (0-100)
    - **rotation**: Watermark rotation angle (0-360)
    - **pages**: Pages to apply watermark: 'all', 'first', 'last', 'custom'
    - **custom_page_range**: Custom page range when pages='custom'
    - **font_size**: Font size in points (default: 36)
    - **color**: Color in hex format (default: #808080 gray)
    - **watermark_image**: Image file for image watermark

    Returns watermarked PDF file for download.
    """
    logger = logging.getLogger(__name__)
    try:
        logger.info(f"Starting watermark processing for file: {file.filename}")
        logger.info(f"Watermark parameters: type={watermark_type}, text='{watermark_text}', position={position}, opacity={opacity}, rotation={rotation}, pages={pages}, font_size={font_size}, color={color}")
        # Validate file
        if not validate_file_type(file, settings.ALLOWED_PDF_TYPES):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File {file.filename} is not a valid PDF. Allowed types: {settings.ALLOWED_PDF_TYPES}"
            )
        
        if not validate_file_size(file):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File {file.filename} exceeds maximum size of {settings.MAX_UPLOAD_SIZE} bytes"
            )
        
        # Validate watermark type
        if watermark_type not in ["text", "image"]:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="watermark_type must be either 'text' or 'image'"
            )
        
        # Validate based on watermark type
        if watermark_type == "text":
            if not watermark_text or len(watermark_text.strip()) == 0:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="watermark_text is required for text watermarks"
                )
        elif watermark_type == "image":
            if not watermark_image:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="watermark_image is required for image watermarks"
                )
            # Validate image file type
            if not validate_file_type(watermark_image, settings.ALLOWED_IMAGE_TYPES):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Image file {watermark_image.filename} is not a valid image. Allowed types: {settings.ALLOWED_IMAGE_TYPES}"
                )
        
        # Validate position
        valid_positions = ["diagonal", "center", "top-left", "top-right", "bottom-left", "bottom-right"]
        if position not in valid_positions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"position must be one of: {', '.join(valid_positions)}"
            )
        
        # Validate opacity
        if opacity < 0 or opacity > 100:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="opacity must be between 0 and 100"
            )
        
        # Validate rotation
        if rotation < 0 or rotation > 360:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="rotation must be between 0 and 360"
            )
        
        # Validate pages
        valid_pages = ["all", "first", "last", "custom"]
        if pages not in valid_pages:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"pages must be one of: {', '.join(valid_pages)}"
            )
        
        # Validate custom page range if pages is custom
        if pages == "custom" and not custom_page_range:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="custom_page_range is required when pages='custom'"
            )
        
        # Validate font size
        if font_size < 8 or font_size > 200:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="font_size must be between 8 and 200 points"
            )
        
        # Validate color format (hex color)
        import re
        color_pattern = re.compile(r'^#([A-Fa-f0-9]{6}|[A-Fa-f0-9]{3})$')
        if not color_pattern.match(color):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="color must be a valid hex color code (e.g., #FF0000 or #F00)"
            )
        
        # Read PDF file
        pdf_bytes = await read_upload_file(file)
        logger.info(f"PDF file read successfully: {len(pdf_bytes)} bytes")
        
        # Read image watermark if provided
        image_bytes = None
        if watermark_type == "image" and watermark_image:
            image_bytes = await read_upload_file(watermark_image)
            logger.info(f"Image watermark read: {len(image_bytes)} bytes")
        
        # Add watermark using enhanced service method
        logger.info(f"Starting watermark processing with parameters: type={watermark_type}, position={position}, opacity={opacity}, rotation={rotation}, pages={pages}, font_size={font_size}, color={color}")
        watermarked_pdf = await run_blocking(
            PDFService.add_watermark_enhanced,
            pdf_bytes=pdf_bytes,
            watermark_type=watermark_type,
            watermark_text=watermark_text.strip() if watermark_text else "",
            position=position,
            opacity=opacity,
            rotation=rotation,
            pages=pages,
            custom_page_range=custom_page_range,
            font_size=font_size,
            color=color,
            image_bytes=image_bytes
        )
        
        logger.info(f"Watermark processing completed successfully. Output size: {len(watermarked_pdf)} bytes")
        
        # Return watermarked PDF
        return create_file_response(
            watermarked_pdf,
            filename=f"watermarked_{file.filename}",
            media_type="application/pdf"
        )
        
    except HTTPException as http_exc:
        logger.warning(f"Validation error in watermark request: {http_exc.detail}")
        raise
    except Exception as e:
        logger.error(f"Error in watermark processing: {str(e)}", exc_info=True)
        return handle_pdf_error(e)


@router.post("/pdf-to-word", summary="Convert PDF to Word document")
async def pdf_to_word(file: UploadFile = File(...)):
    """
    Convert a PDF file to an editable Word document (.docx) with high quality.
    
    Uses pdf2docx library for better preservation of:
    - Paragraph formatting
    - Tables and structure
    - Text layout
    - Multi-page content
    
    Falls back to OCR for scanned/image-based PDFs.
    
    - **file**: PDF file to convert
    
    Returns converted Word document for download.
    """
    import os
    import uuid
    import tempfile
    import traceback
    from pathlib import Path
    
    conversion_logger = logging.getLogger("pdf_to_word")
    conversion_logger.setLevel(logging.DEBUG)
    
    try:
        conversion_logger.info("=" * 60)
        conversion_logger.info(f"PDF-to-Word conversion STARTED: {file.filename}")
        conversion_logger.info(f"Content-Type: {file.content_type}")
        
        # --- STEP 1: Validate file type ---
        if not file.filename:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No filename provided. Please upload a valid PDF file."
            )
        
        if not validate_file_type(file, settings.ALLOWED_PDF_TYPES):
            conversion_logger.warning(f"File type rejected: {file.content_type} for {file.filename}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Invalid file type. Please upload a PDF file (received: {file.content_type or 'unknown'})."
            )
        
        # --- STEP 2: Validate file size ---
        if not validate_file_size(file):
            max_mb = settings.MAX_UPLOAD_SIZE // (1024 * 1024)
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File is too large. Maximum allowed size is {max_mb} MB."
            )
        
        # --- STEP 3: Read file ---
        pdf_bytes = await read_upload_file(file)
        conversion_logger.info(f"File read successfully: {len(pdf_bytes):,} bytes ({len(pdf_bytes) / 1024:.1f} KB)")
        
        if len(pdf_bytes) == 0:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="The uploaded file is empty. Please upload a valid PDF file."
            )
        
        # --- STEP 4: Validate PDF integrity ---
        integrity_ok, integrity_msg = await run_blocking(
            _validate_pdf_integrity, pdf_bytes, file.filename or "unknown"
        )
        if not integrity_ok:
            conversion_logger.warning(f"PDF integrity check FAILED: {integrity_msg}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"The PDF file appears to be corrupted or invalid: {integrity_msg}"
            )
        conversion_logger.info(f"PDF integrity check PASSED: {integrity_msg}")

        # --- STEP 5: Detect if PDF is scanned/image-based ---
        is_scanned, scan_info = await run_blocking(_detect_scanned_pdf, pdf_bytes)
        conversion_logger.info(f"Scanned PDF detection: is_scanned={is_scanned}, info={scan_info}")
        
        # --- STEP 6: Convert in temp directory with UUID-based filenames ---
        with tempfile.TemporaryDirectory(prefix="pdf2word_") as temp_dir:
            temp_dir_path = Path(temp_dir)
            
            # Use UUID-based filenames to avoid collisions and stale files
            unique_id = uuid.uuid4().hex[:12]
            pdf_filename = f"input_{unique_id}.pdf"
            docx_filename = f"output_{unique_id}.docx"
            
            pdf_path = temp_dir_path / pdf_filename
            docx_path = temp_dir_path / docx_filename
            
            conversion_logger.info(f"Temp dir: {temp_dir_path}")
            conversion_logger.info(f"PDF path: {pdf_path}")
            conversion_logger.info(f"DOCX path: {docx_path}")
            
            # Write PDF to temp file
            with open(pdf_path, "wb") as f:
                f.write(pdf_bytes)
            conversion_logger.info(f"PDF written to temp file: {pdf_path} ({pdf_path.stat().st_size:,} bytes)")
            
            # --- STEP 6a: If scanned PDF, try OCR first ---
            if is_scanned:
                conversion_logger.info(f"PDF appears to be scanned/image-based. Attempting OCR conversion...")
                try:
                    async with heavy_job_slot():
                        ocr_docx = await run_blocking(
                            _convert_with_ocr, pdf_bytes, file.filename, pdf_path, docx_path
                        )
                    if ocr_docx and len(ocr_docx) > 1000:
                        conversion_logger.info(f"OCR conversion succeeded: {len(ocr_docx):,} bytes")
                        output_filename = f"{Path(file.filename).stem}_converted.docx"
                        return create_file_response(
                            ocr_docx,
                            filename=output_filename,
                            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        conversion_logger.warning(f"OCR produced insufficient output ({len(ocr_docx) if ocr_docx else 0} bytes). Falling through to normal conversion.")
                except Exception as ocr_error:
                    conversion_logger.warning(f"OCR conversion failed, falling through to normal methods: {ocr_error}")
            
            # --- STEP 6b: Primary - pdf2docx (works on all platforms with multi_processing=False) ---
            try:
                from pdf2docx import Converter

                conversion_logger.info("Attempting pdf2docx conversion (multi_processing=False)...")

                # pdf2docx is fully blocking/CPU-bound. Run the whole
                # create→convert→close cycle in one thread so the Converter
                # object never crosses thread boundaries, keeping the event
                # loop free for other requests.
                def _run_pdf2docx(src: str, dst: str):
                    cv = Converter(src)
                    try:
                        cv.convert(dst, start=0, end=None, multi_processing=False, debug=False)
                    finally:
                        cv.close()

                async with heavy_job_slot():
                    await run_blocking(_run_pdf2docx, str(pdf_path), str(docx_path))
                
                conversion_logger.info("pdf2docx Converter finished.")
                
                if not docx_path.exists():
                    conversion_logger.error("pdf2docx did NOT produce an output file")
                    raise PDFProcessingError("No output file generated by pdf2docx")
                
                docx_size = docx_path.stat().st_size
                conversion_logger.info(f"pdf2docx output file size: {docx_size:,} bytes")
                
                if docx_size == 0:
                    conversion_logger.error("pdf2docx produced an empty output file")
                    raise PDFProcessingError("Empty output file from pdf2docx")
                
                # Read the DOCX
                with open(docx_path, "rb") as f:
                    docx_bytes = f.read()
                
                # Verify the DOCX is a valid ZIP (DOCX files are ZIP archives)
                import zipfile
                try:
                    with zipfile.ZipFile(docx_path, 'r') as zipf:
                        image_files = [f for f in zipf.namelist() if 'media' in f and f.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'))]
                        conversion_logger.info(f"DOCX integrity OK. Contains {len(image_files)} embedded images.")
                        for img in image_files:
                            conversion_logger.debug(f"  Image: {img}")
                except zipfile.BadZipFile:
                    conversion_logger.error("pdf2docx produced an invalid DOCX (not a valid ZIP)")
                    raise PDFProcessingError("pdf2docx produced a corrupt DOCX file")
                
                output_filename = f"{Path(file.filename).stem}_converted.docx"
                conversion_logger.info(f"pdf2docx conversion SUCCESS: {output_filename} ({len(docx_bytes):,} bytes)")
                conversion_logger.info(f"Response MIME type: application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                conversion_logger.info(f"Response filename: {output_filename}")
                conversion_logger.info(f"Content-Disposition: attachment; filename=\"{output_filename}\"")
                
                return create_file_response(
                    docx_bytes,
                    filename=output_filename,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as pdf2docx_err:
                conversion_logger.warning(f"pdf2docx failed: {type(pdf2docx_err).__name__}: {pdf2docx_err}")
                conversion_logger.debug(f"pdf2docx traceback: {traceback.format_exc()}")
            
            # --- STEP 6c: Fallback - pdfplumber + python-docx ---
            try:
                conversion_logger.info("Attempting pdfplumber + python-docx conversion...")
                async with heavy_job_slot():
                    alt_docx = await run_blocking(convert_with_pdfplumber, pdf_bytes, file.filename)
                
                if alt_docx and len(alt_docx) > 500:
                    conversion_logger.info(f"pdfplumber conversion SUCCESS: {len(alt_docx):,} bytes")
                    return create_file_response(
                        alt_docx,
                        filename=f"{Path(file.filename).stem}_converted.docx",
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    conversion_logger.warning(f"pdfplumber produced insufficient output ({len(alt_docx) if alt_docx else 0} bytes)")
                    raise PDFProcessingError("pdfplumber produced insufficient output")
                    
            except Exception as pdfplumber_err:
                conversion_logger.warning(f"pdfplumber failed: {type(pdfplumber_err).__name__}: {pdfplumber_err}")
            
            # --- STEP 6d: Final fallback - basic text extraction ---
            try:
                conversion_logger.info("Attempting basic text extraction (final fallback)...")
                basic_docx = await run_blocking(create_basic_word_document, pdf_bytes, file.filename)
                
                if basic_docx and len(basic_docx) > 200:
                    conversion_logger.info(f"Basic extraction SUCCESS: {len(basic_docx):,} bytes")
                    return create_file_response(
                        basic_docx,
                        filename=f"{Path(file.filename).stem}_converted.docx",
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    raise PDFProcessingError("All conversion methods produced insufficient output")
                    
            except Exception as basic_err:
                conversion_logger.error(f"ALL conversion methods FAILED: {basic_err}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="Could not convert this PDF to Word. The file may be heavily corrupted or use an unsupported format. Please try a different PDF."
                )
        
    except HTTPException:
        raise
    except PDFProcessingError as e:
        conversion_logger.error(f"PDFProcessingError: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"PDF conversion failed: {str(e)}. Please try a different PDF file."
        )
    except Exception as e:
        conversion_logger.error(f"Unexpected error: {type(e).__name__}: {e}")
        conversion_logger.error(traceback.format_exc())
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="An unexpected error occurred during PDF conversion. Please try again later."
        )


def _validate_pdf_integrity(pdf_bytes: bytes, filename: str) -> tuple:
    """
    Validate PDF file integrity by checking the PDF header and structure.
    Returns (is_valid: bool, message: str).
    """
    import re
    
    if len(pdf_bytes) < 10:
        return False, "File is too small to be a valid PDF"
    
    # Check PDF header (must start with %PDF-)
    header = pdf_bytes[:8]
    if not header.startswith(b'%PDF-'):
        # Check for UTF-8 BOM
        if header.startswith(b'\xef\xbb\xbf%PDF-'):
            pass  # Valid with BOM
        else:
            return False, f"Missing PDF header. File starts with: {header[:20]!r}"
    
    # Check for PDF version
    try:
        header_str = pdf_bytes[:1024].decode('latin-1', errors='ignore')
        version_match = re.search(r'%PDF-(\d+\.\d+)', header_str)
        if version_match:
            pdf_version = version_match.group(1)
        else:
            return False, "Could not determine PDF version"
    except Exception:
        pdf_version = "unknown"
    
    # Check for %%EOF marker near the end
    trailer = pdf_bytes[-1024:]
    if b'%%EOF' not in trailer:
        # File might be truncated but could still be partially readable
        pass
    
    # Try basic PyPDF2 read to check if parseable
    try:
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        num_pages = len(pdf_reader.pages)
        return True, f"PDF v{pdf_version}, {num_pages} page(s), {len(pdf_bytes):,} bytes"
    except Exception as e:
        # PyPDF2 failed - might still be partially recoverable
        return True, f"PDF v{pdf_version}, {len(pdf_bytes):,} bytes (warning: {str(e)[:100]})"


def _detect_scanned_pdf(pdf_bytes: bytes) -> tuple:
    """
    Detect if a PDF is scanned (image-based) by checking for extractable text.
    Returns (is_scanned: bool, info: str).
    """
    try:
        import PyPDF2
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        num_pages = len(pdf_reader.pages)
        
        if num_pages == 0:
            return False, "0 pages"
        
        # Sample up to first 3 pages
        pages_to_check = min(3, num_pages)
        total_text = ""
        
        for i in range(pages_to_check):
            try:
                page_text = pdf_reader.pages[i].extract_text()
                if page_text:
                    total_text += page_text
            except Exception:
                continue
        
        # Heuristic: if total extracted text across sampled pages is very small,
        # the PDF is likely scanned/image-based
        avg_text_per_page = len(total_text.strip()) / pages_to_check if pages_to_check > 0 else 0
        
        if avg_text_per_page < 20:
            return True, f"Scanned/image-based PDF detected ({num_pages} pages, avg {avg_text_per_page:.0f} chars/page)"
        else:
            return False, f"Text-based PDF ({num_pages} pages, avg {avg_text_per_page:.0f} chars/page)"
            
    except Exception as e:
        return False, f"Detection skipped: {type(e).__name__}"


def _convert_with_ocr(pdf_bytes: bytes, original_filename: str, pdf_path, docx_path) -> bytes:
    """
    Convert a scanned/image-based PDF to Word using OCR (pytesseract).
    Returns the DOCX bytes or raises an exception.
    """
    import logging
    ocr_logger = logging.getLogger("pdf_to_word.ocr")
    
    ocr_logger.info("Starting OCR-based PDF conversion...")
    
    try:
        from pdf2image import convert_from_bytes
        from docx import Document
        from docx.shared import Inches, Pt
        import pytesseract
        
        ocr_logger.info("pdf2image + pytesseract imported successfully")
        
        # Convert PDF pages to images
        ocr_logger.info("Converting PDF pages to images...")
        try:
            images = convert_from_bytes(pdf_bytes, dpi=300)
        except Exception as img_err:
            ocr_logger.warning(f"pdf2image at 300 DPI failed: {img_err}. Trying 150 DPI...")
            try:
                images = convert_from_bytes(pdf_bytes, dpi=150)
            except Exception:
                raise PDFProcessingError(f"Could not render PDF pages for OCR: {img_err}")
        
        ocr_logger.info(f"Converted {len(images)} page(s) to images at {images[0].size if images else 'unknown'} resolution")
        
        if not images:
            raise PDFProcessingError("No pages could be extracted from the PDF for OCR")
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'OCR Conversion: {original_filename}', 0)
        
        # Add metadata note
        meta = doc.add_paragraph()
        meta.add_run('This document was converted from a scanned PDF using OCR (Optical Character Recognition). ').italic = True
        meta.add_run('Some formatting may differ from the original.').italic = True
        doc.add_paragraph()
        
        # OCR each page
        for page_num, image in enumerate(images, 1):
            ocr_logger.debug(f"OCR processing page {page_num}/{len(images)}...")
            
            # Add page header
            doc.add_heading(f'Page {page_num}', level=1)
            
            # Perform OCR
            try:
                page_text = pytesseract.image_to_string(image, lang='eng')
            except Exception as tesseract_err:
                ocr_logger.warning(f"pytesseract failed for page {page_num}: {tesseract_err}")
                page_text = f"[OCR failed for page {page_num}: {tesseract_err}]"
            
            if page_text and page_text.strip():
                # Split into paragraphs on double newlines
                paragraphs = page_text.split('\n\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        p = doc.add_paragraph(para_text)
                        p.style.font.size = Pt(11)
            else:
                doc.add_paragraph('[No text detected on this page. The image may be blank or the quality too low for OCR.]')
            
            # Page break between pages
            if page_num < len(images):
                doc.add_page_break()
        
        # Add footer note
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run('Generated using OCR technology. ').italic = True
        footer.add_run('For best results, ensure the original PDF has clear, high-quality text.').italic = True
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        result = buffer.getvalue()
        
        ocr_logger.info(f"OCR conversion completed: {len(result):,} bytes")
        return result
        
    except ImportError as ie:
        ocr_logger.warning(f"OCR dependencies not available: {ie}")
        raise PDFProcessingError(f"OCR is not available (missing dependencies: {ie}). Please install pytesseract and pdf2image.")
    except PDFProcessingError:
        raise
    except Exception as e:
        ocr_logger.error(f"OCR conversion failed: {type(e).__name__}: {e}")
        ocr_logger.debug(traceback.format_exc())
        raise PDFProcessingError(f"OCR conversion failed: {str(e)[:200]}")


def create_simulated_word_document(pdf_bytes: bytes, original_filename: str) -> bytes:
    """
    Create a Word document with actual PDF content extracted as text.
    Returns a valid .docx file that can be opened in Word.
    """
    import zipfile
    import io
    from datetime import datetime
    import PyPDF2
    
    # Extract text from PDF
    extracted_text = ""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        num_pages = len(pdf_reader.pages)
        
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            text = page.extract_text()
            if text:
                extracted_text += f"--- Page {page_num + 1} ---\n{text}\n\n"
        
        if not extracted_text:
            extracted_text = "[No text could be extracted from the PDF. The PDF may contain only images or scanned content.]"
            
    except Exception as e:
        extracted_text = f"[Error extracting text from PDF: {str(e)}]"
    
    # Limit text size to avoid XML issues (Word has limits)
    if len(extracted_text) > 50000:
        extracted_text = extracted_text[:50000] + "\n\n[Content truncated due to size limits...]"
    
    # Escape XML special characters
    import html
    extracted_text_escaped = html.escape(extracted_text)
    
    # Create a minimal valid .docx file (Office Open XML format)
    # A .docx file is a ZIP archive containing XML files
    
    buffer = io.BytesIO()
    
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
        # Create minimal required files for a .docx
        
        # [Content_Types].xml
        content_types = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/_rels/document.xml.rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
</Types>'''
        zipf.writestr('[Content_Types].xml', content_types)
        
        # _rels/.rels
        rels = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>'''
        zipf.writestr('_rels/.rels', rels)
        
        # word/_rels/document.xml.rels
        doc_rels = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
</Relationships>'''
        zipf.writestr('word/_rels/document.xml.rels', doc_rels)
        
        # word/document.xml - the actual document content with extracted PDF text
        now = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
        doc_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
            xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">
  <w:body>
    <w:p>
      <w:r>
        <w:t>PDF to Word Conversion</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:r>
        <w:t>File: {html.escape(original_filename)}</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:r>
        <w:t>Converted on: {now}</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:r>
        <w:t>Original PDF size: {len(pdf_bytes)} bytes</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:r>
        <w:t> </w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:r>
        <w:t>EXTRACTED PDF CONTENT:</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:r>
        <w:t>{extracted_text_escaped}</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:r>
        <w:t> </w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:r>
        <w:t>Note: This document contains text extracted from the PDF. For full formatting preservation,</w:t>
      </w:r>
    </w:p>
    <w:p>
      <w:r>
        <w:t>ensure LibreOffice is properly installed with DOCX export support.</w:t>
      </w:r>
    </w:p>
    <w:sectPr>
      <w:pgSz w:w="12240" w:h="15840"/>
      <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="720" w:footer="720" w:gutter="0"/>
      <w:cols w:space="720"/>
    </w:sectPr>
  </w:body>
</w:document>'''
        zipf.writestr('word/document.xml', doc_xml)
    
    buffer.seek(0)
    return buffer.getvalue()


def create_improved_word_document(pdf_bytes: bytes, original_filename: str) -> bytes:
    """
    Create an improved Word document with better formatting and structure.
    Uses pdfplumber for text extraction with layout information and python-docx for document creation.
    """
    try:
        import pdfplumber
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        from datetime import datetime
        
        # Create a new Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'PDF to Word Conversion: {original_filename}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        metadata = doc.add_paragraph()
        metadata.add_run(f'Original PDF: {original_filename}\n')
        metadata.add_run(f'Conversion Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
        metadata.add_run(f'File Size: {len(pdf_bytes):,} bytes\n')
        
        doc.add_paragraph()  # Add spacing
        
        # Extract text from PDF with pdfplumber (preserves layout better)
        try:
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Add page heading
                    page_heading = doc.add_heading(f'Page {page_num}', level=1)
                    
                    # Extract text with layout
                    text = page.extract_text()
                    if text:
                        # Add text to document with basic formatting
                        para = doc.add_paragraph()
                        run = para.add_run(text)
                        run.font.size = Pt(11)
                    
                    # Try to extract tables
                    tables = page.extract_tables()
                    if tables:
                        table_heading = doc.add_heading(f'Tables on Page {page_num}', level=2)
                        
                        for table_num, table in enumerate(tables, 1):
                            if table and any(any(cell for cell in row) for row in table):
                                # Create Word table
                                word_table = doc.add_table(rows=len(table), cols=len(table[0]) if table else 1)
                                word_table.style = 'Light Grid Accent 1'
                                
                                # Populate table cells
                                for i, row in enumerate(table):
                                    for j, cell in enumerate(row):
                                        if cell:
                                            word_table.cell(i, j).text = str(cell)
                                
                                doc.add_paragraph()  # Add spacing after table
                    
                    # Add page break between pages (except last page)
                    if page_num < len(pdf.pages):
                        doc.add_page_break()
                        
        except Exception as pdf_error:
            # Fallback to basic text extraction if pdfplumber fails
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    doc.add_heading(f'Page {page_num + 1}', level=1)
                    doc.add_paragraph(text)
                if page_num < len(pdf_reader.pages) - 1:
                    doc.add_page_break()
        
        # Add footer note
        doc.add_paragraph()
        note = doc.add_paragraph('Note: This document was generated using improved PDF extraction with layout preservation. ')
        note.add_run('Tables and formatting are preserved where possible.')
        note.runs[0].italic = True
        
        # Save document to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        # If anything fails, fall back to basic simulated document
        logger = logging.getLogger(__name__)
        logger.warning(f"Improved Word document creation failed: {e}, falling back to basic")
        return create_basic_word_document(pdf_bytes, original_filename)


def convert_with_pdfplumber(pdf_bytes: bytes, original_filename: str) -> bytes:
    """
    Alternative conversion using pdfplumber for better text extraction with layout.
    Creates a Word document with preserved paragraph structure.
    """
    try:
        import pdfplumber
        from docx import Document
        from docx.shared import Pt
        import io
        
        doc = Document()
        
        # Add header
        doc.add_heading(f'PDF to Word: {original_filename}', 0)
        
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text with layout using pdfplumber's improved extraction
                text = page.extract_text(
                    x_tolerance=3,  # Horizontal tolerance for joining words
                    y_tolerance=3,  # Vertical tolerance for joining lines
                    layout=False,    # Don't use layout analysis (faster)
                    x_density=7.25,  # Characters per inch horizontally
                    y_density=13     # Lines per inch vertically
                )
                
                if text:
                    # Split text into paragraphs based on double newlines
                    paragraphs = text.split('\n\n')
                    
                    # Add page heading
                    doc.add_heading(f'Page {page_num}', level=1)
                    
                    # Add each paragraph
                    for para_text in paragraphs:
                        if para_text.strip():
                            paragraph = doc.add_paragraph(para_text.strip())
                            paragraph.style.font.size = Pt(11)
                
                # Try to extract and preserve tables
                tables = page.extract_tables()
                if tables:
                    for table_num, table in enumerate(tables, 1):
                        if table and any(any(cell for cell in row) for row in table):
                            doc.add_heading(f'Table {table_num}', level=2)
                            
                            # Create Word table
                            word_table = doc.add_table(
                                rows=len(table),
                                cols=max(len(row) for row in table) if table else 1
                            )
                            
                            # Populate table
                            for i, row in enumerate(table):
                                for j, cell in enumerate(row):
                                    if cell:
                                        word_table.cell(i, j).text = str(cell)
                            
                            doc.add_paragraph()  # Spacing after table
                
                # Page break for next page
                if page_num < len(pdf.pages):
                    doc.add_page_break()
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        logger = logging.getLogger(__name__)
        logger.error(f"pdfplumber conversion failed: {e}")
        # Fall back to basic document
        return create_basic_word_document(pdf_bytes, original_filename)


def create_basic_word_document(pdf_bytes: bytes, original_filename: str) -> bytes:
    """
    Create a basic Word document with text extraction only.
    This is the final fallback when other methods fail.
    """
    try:
        import PyPDF2
        from docx import Document
        from docx.shared import Pt
        import io
        from datetime import datetime
        
        doc = Document()
        
        # Add title
        doc.add_heading('PDF to Word Conversion', 0)
        
        # Add metadata
        doc.add_paragraph(f'Original PDF: {original_filename}')
        doc.add_paragraph(f'Converted: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'File size: {len(pdf_bytes):,} bytes')
        doc.add_paragraph()
        
        # Extract text from PDF
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        all_text = ""
        
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text = page.extract_text()
            if text:
                all_text += f"=== Page {page_num + 1} ===\n{text}\n\n"
        
        if not all_text:
            all_text = "No text could be extracted from the PDF. The file may contain only images or scanned content."
        
        # Add extracted text
        doc.add_heading('Extracted Content', level=1)
        content_para = doc.add_paragraph(all_text)
        
        # Set font size
        for run in content_para.runs:
            run.font.size = Pt(10)
        
        # Add note
        doc.add_paragraph()
        note = doc.add_paragraph('Note: This is a basic text extraction. For better formatting preservation, ')
        note.add_run('ensure all required PDF processing libraries are installed.')
        note.runs[1].italic = True
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        # Ultimate fallback to the original simulated document
        logger = logging.getLogger(__name__)
        logger.error(f"Basic Word document creation also failed: {e}, using simulated document")
        return create_simulated_word_document(pdf_bytes, original_filename)


@router.post("/word-to-pdf", summary="Convert Word document to PDF")
async def word_to_pdf(file: UploadFile = File(...)):
    """
    Convert Word document (.doc, .docx) to PDF format using LibreOffice.
    
    Parameters:
    - file: Word document file to convert
    
    Returns converted PDF file.
    """
    import tempfile
    import subprocess
    import os
    from pathlib import Path
    
    logger.debug(f"Received Word file: {file.filename}")
    
    try:
        # Validate file type
        if not file.filename:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="File must have a filename"
            )
        
        # Check if it's a Word document
        filename_lower = file.filename.lower()
        if not (filename_lower.endswith('.doc') or filename_lower.endswith('.docx')):
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail="File must be a Word document (.doc or .docx)"
            )
        
        # Read file content
        file_bytes = await read_upload_file(file)
        logger.debug(f"File size: {len(file_bytes)} bytes")
        
        # Create temporary directory (auto-cleanup)
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_dir_path = Path(temp_dir)
            logger.debug(f"Created temp directory: {temp_dir}")
            
            # Save uploaded file
            input_path = temp_dir_path / file.filename
            with open(input_path, "wb") as f:
                f.write(file_bytes)
            logger.debug(f"Saved uploaded file to: {input_path}")
            
            # Convert using LibreOffice (path auto-resolved across OSes)
            libreoffice_path = resolve_libreoffice_path()

            if not libreoffice_path:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="LibreOffice is not installed on the server"
                )

            logger.debug("Starting LibreOffice conversion...")
            # Unique per-request profile dir so parallel LibreOffice runs don't
            # collide on the shared default user profile lock.
            import uuid as _uuid
            profile_dir = temp_dir_path / f"lo_profile_{_uuid.uuid4().hex[:8]}"
            cmd = [
                libreoffice_path,
                '--headless',
                f'-env:UserInstallation=file:///{profile_dir.as_posix().lstrip("/")}',
                '--convert-to', 'pdf',
                '--outdir', str(temp_dir_path),
                str(input_path)
            ]

            try:
                # Run the blocking subprocess off the event loop, and bound how
                # many heavy LibreOffice jobs run at once.
                async with heavy_job_slot():
                    result = await run_blocking(
                        subprocess.run,
                        cmd,
                        capture_output=True,
                        text=True,
                        timeout=60,  # 60 second timeout
                    )
                
                logger.debug(f"LibreOffice return code: {result.returncode}")
                logger.debug(f"LibreOffice stdout: {result.stdout}")
                
                if result.returncode != 0:
                    logger.debug(f"LibreOffice stderr: {result.stderr}")
                    raise HTTPException(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        detail=f"PDF conversion failed: {result.stderr[:200]}"
                    )
                
            except subprocess.TimeoutExpired:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="PDF conversion timed out"
                )
            
            # Find converted PDF
            pdf_files = list(temp_dir_path.glob('*.pdf'))
            if not pdf_files:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="PDF conversion completed but no PDF file was generated"
                )
            
            pdf_file = pdf_files[0]
            logger.debug(f"Converted PDF: {pdf_file}")
            
            # Read PDF file
            with open(pdf_file, "rb") as f:
                pdf_bytes = f.read()
            
            logger.debug(f"PDF size: {len(pdf_bytes)} bytes")
            
            # Generate output filename
            original_name = Path(file.filename).stem
            output_filename = f"{original_name}_converted.pdf"
            
            logger.debug(f"Returning PDF: {output_filename}")
            logger.debug("Cleaning up temporary files...")
            
            # Return PDF file (temp files auto-deleted when context exits)
            return create_file_response(
                pdf_bytes,
                output_filename,
                media_type="application/pdf"
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger = logging.getLogger(__name__)
        logger.error(f"Word to PDF conversion failed: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to convert Word to PDF: {str(e)}"
        )


@router.post("/compress-pdf", summary="Compress PDF file")
async def compress_pdf(
    file: UploadFile = File(...),
    compressionLevel: str = Form("medium")
):
    """
    Compress PDF file to reduce file size.

    Parameters:
    - file: PDF file to compress
    - compressionLevel: Compression level (low, medium, high)

    Returns compressed PDF file.
    """
    try:
        # Validate file type
        if not file.content_type or "pdf" not in file.content_type.lower():
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail="File must be a PDF"
            )

        # Validate compression level
        if compressionLevel not in ["low", "medium", "high"]:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="compressionLevel must be 'low', 'medium', or 'high'"
            )

        # Read file content
        file_bytes = await read_upload_file(file)
        original_filename = file.filename or "document.pdf"
        original_size = len(file_bytes)

        logger.info(
            "Compressing '%s' (%d bytes) with %s compression",
            original_filename, original_size, compressionLevel
        )

        # Delegate to the service layer (uses temp files, handles latin-1 fallback)
        compressed_bytes = await run_blocking(PDFService.compress_pdf, file_bytes, compressionLevel)

        compressed_size = len(compressed_bytes)
        reduction_pct = (1 - compressed_size / max(original_size, 1)) * 100
        logger.info(
            "Compression complete: %d → %d bytes (%.1f%% reduction)",
            original_size, compressed_size, reduction_pct
        )

        # Generate output filename
        output_filename = "pdforca-compressfile.pdf"

        # Return compressed file
        return create_file_response(
            compressed_bytes,
            output_filename,
            media_type="application/pdf"
        )

    except PDFProcessingError as e:
        logger.error("PDF compression service error: %s", e)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e)
        )
    except ValueError as e:
        logger.error("PDF compression validation error: %s", e)
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=str(e)
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Unexpected PDF compression error: %s", e, exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to compress PDF: {str(e)}"
        )


@router.post("/fix-scanned-pdf", summary="Fix scanned PDF documents")
async def fix_scanned_pdf(file: UploadFile = File(...)):
    """
    Fix scanned PDF documents by optimizing for OCR and readability.
    
    Parameters:
    - file: Scanned PDF file to fix
    
    Returns optimized PDF file.
    """
    logger = logging.getLogger(__name__)
    try:
        # Validate file type
        if not file.content_type or "pdf" not in file.content_type.lower():
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail="File must be a PDF"
            )
        
        # Read file content
        file_bytes = await read_upload_file(file)
        
        # Process PDF using PDFOptimizationService
        from app.services.pdf_service import PDFOptimizationService
        processed_bytes = await run_blocking(PDFOptimizationService.fix_scanned_pdf, file_bytes)

        # Create filename
        original_filename = file.filename or "document.pdf"
        output_filename = f"fixed_{original_filename}"
        
        # Return processed PDF
        return create_file_response(
            file_bytes=processed_bytes,
            filename=output_filename,
            media_type="application/pdf"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to fix scanned PDF: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to fix scanned PDF: {str(e)}"
        )


@router.post("/optimize-pdf", summary="Optimize PDF for web viewing")
async def optimize_pdf(file: UploadFile = File(...)):
    """
    Optimize PDF for web viewing with linearization and compression.
    
    Parameters:
    - file: PDF file to optimize
    
    Returns optimized PDF file.
    """
    logger = logging.getLogger(__name__)
    try:
        # Validate file type
        if not file.content_type or "pdf" not in file.content_type.lower():
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail="File must be a PDF"
            )
        
        # Read file content
        file_bytes = await read_upload_file(file)
        
        # Process PDF using PDFOptimizationService
        from app.services.pdf_service import PDFOptimizationService
        processed_bytes = await run_blocking(PDFOptimizationService.optimize_for_viewing, file_bytes)
        
        # Create filename
        original_filename = file.filename or "document.pdf"
        output_filename = f"optimized_{original_filename}"
        
        # Return processed PDF
        return create_file_response(
            file_bytes=processed_bytes,
            filename=output_filename,
            media_type="application/pdf"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to optimize PDF: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to optimize PDF: {str(e)}"
        )


@router.post("/prepare-print-pdf", summary="Prepare PDF for printing")
async def prepare_print_pdf(file: UploadFile = File(...)):
    """
    Prepare PDF for printing with proper page sizing and margins.
    
    Parameters:
    - file: PDF file to prepare for printing
    
    Returns print-ready PDF file.
    """
    logger = logging.getLogger(__name__)
    try:
        # Validate file type
        if not file.content_type or "pdf" not in file.content_type.lower():
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail="File must be a PDF"
            )
        
        # Read file content
        file_bytes = await read_upload_file(file)
        
        # Process PDF using PDFOptimizationService
        from app.services.pdf_service import PDFOptimizationService
        processed_bytes = await run_blocking(PDFOptimizationService.prepare_for_printing, file_bytes)
        
        # Create filename
        original_filename = file.filename or "document.pdf"
        output_filename = f"print_ready_{original_filename}"
        
        # Return processed PDF
        return create_file_response(
            file_bytes=processed_bytes,
            filename=output_filename,
            media_type="application/pdf"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to prepare PDF for printing: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to prepare PDF for printing: {str(e)}"
        )


@router.post("/protect-pdf", summary="Protect PDF with password and encryption")
async def protect_pdf(
    file: UploadFile = File(...),
    password: str = Form(..., min_length=4, description="Password to protect the PDF"),
    allow_printing: bool = Form(True, description="Allow printing"),
    allow_copying: bool = Form(True, description="Allow copying text/images"),
    allow_editing: bool = Form(True, description="Allow modifying content"),
    allow_annotating: bool = Form(True, description="Allow adding annotations"),
):
    """
    Protect a PDF with password encryption using AES-256.
    Returns the encrypted PDF file for download.
    """
    try:
        # Read the uploaded PDF
        pdf_bytes = await file.read()

        # Validate the file is a PDF
        if not file.filename or not file.filename.lower().endswith(".pdf"):
            raise HTTPException(
                status_code=400,
                detail="Invalid file type. Please upload a PDF file.",
            )

        if len(pdf_bytes) == 0:
            raise HTTPException(
                status_code=400,
                detail="The uploaded PDF file is empty.",
            )

        # Validate password
        if len(password) < 4:
            raise HTTPException(
                status_code=400,
                detail="Password must be at least 4 characters long.",
            )

        logger.debug(f"[protect_pdf] Encrypting PDF: {file.filename} "
                     f"(printing={allow_printing}, copying={allow_copying}, "
                     f"editing={allow_editing}, annotating={allow_annotating})")

        # Encrypt the PDF
        encrypted_pdf = await run_blocking(
            PDFService.protect_pdf,
            pdf_bytes=pdf_bytes,
            password=password,
            allow_printing=allow_printing,
            allow_copying=allow_copying,
            allow_editing=allow_editing,
            allow_annotating=allow_annotating,
        )

        # Generate output filename
        base_name = os.path.splitext(file.filename)[0]
        output_filename = f"{base_name}_protected.pdf"

        logger.debug(f"[protect_pdf] Successfully encrypted PDF, returning "
                     f"{len(encrypted_pdf)} bytes as '{output_filename}'")

        return StreamingResponse(
            io.BytesIO(encrypted_pdf),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{output_filename}"',
                "Content-Length": str(len(encrypted_pdf)),
                "X-Protected-PDF": "true",
            },
        )

    except HTTPException:
        raise
    except PDFProcessingError as e:
        logger.debug(f"[protect_pdf] PDFProcessingError: {e.message}")
        raise HTTPException(
            status_code=500,
            detail=f"PDF protection failed: {e.message}",
        )
    except Exception as e:
        logger.debug(f"[protect_pdf] Unexpected error: {type(e).__name__}: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"An unexpected error occurred while protecting the PDF.",
        )


@router.post("/unlock-pdf", summary="Remove password from a PDF")
async def unlock_pdf(
    file: UploadFile = File(...),
    password: str = Form(..., description="Password used to open the PDF"),
):
    """Decrypt an AES/RC4-encrypted PDF and return the unlocked file."""
    try:
        pdf_bytes = await file.read()
        if not file.filename or not file.filename.lower().endswith(".pdf"):
            raise HTTPException(status_code=400, detail="Invalid file type. Please upload a PDF file.")
        if len(pdf_bytes) == 0:
            raise HTTPException(status_code=400, detail="The uploaded PDF file is empty.")

        unlocked = await run_blocking(PDFService.unlock_pdf, pdf_bytes=pdf_bytes, password=password)
        base = os.path.splitext(file.filename)[0]
        output_filename = f"{base}_unlocked.pdf"
        return StreamingResponse(
            io.BytesIO(unlocked),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{output_filename}"',
                "Content-Length": str(len(unlocked)),
            },
        )
    except HTTPException:
        raise
    except PDFProcessingError as e:
        # Wrong password is a 400, not a 500
        status_code = 400 if "password" in e.message.lower() else 500
        raise HTTPException(status_code=status_code, detail=e.message)
    except Exception as e:
        logger.exception("[unlock_pdf] unexpected error")
        raise HTTPException(status_code=500, detail="An unexpected error occurred while unlocking the PDF.")


@router.post("/rotate-pdf", summary="Rotate pages in a PDF")
async def rotate_pdf(
    file: UploadFile = File(...),
    angle: int = Form(90, description="Rotation angle: 90, 180, or 270"),
    page_range: str = Form("all", description='Pages to rotate: "all" or "1,3,5-7"'),
):
    """Rotate selected pages of a PDF by 90/180/270 degrees."""
    try:
        pdf_bytes = await file.read()
        if not file.filename or not file.filename.lower().endswith(".pdf"):
            raise HTTPException(status_code=400, detail="Invalid file type. Please upload a PDF file.")
        if len(pdf_bytes) == 0:
            raise HTTPException(status_code=400, detail="The uploaded PDF file is empty.")

        rotated = await run_blocking(PDFService.rotate_pdf, pdf_bytes=pdf_bytes, angle=angle, page_range=page_range)
        base = os.path.splitext(file.filename)[0]
        output_filename = f"{base}_rotated.pdf"
        return StreamingResponse(
            io.BytesIO(rotated),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{output_filename}"',
                "Content-Length": str(len(rotated)),
            },
        )
    except HTTPException:
        raise
    except PDFProcessingError as e:
        raise HTTPException(status_code=400, detail=e.message)
    except Exception as e:
        logger.exception("[rotate_pdf] unexpected error")
        raise HTTPException(status_code=500, detail="An unexpected error occurred while rotating the PDF.")


@router.post("/extract-pages", summary="Extract selected pages into a new PDF")
async def extract_pages(
    file: UploadFile = File(...),
    pages: str = Form(..., description='Pages to extract, e.g. "1,3,5-7"'),
):
    """Pull selected pages from a PDF into a single new PDF."""
    try:
        pdf_bytes = await file.read()
        if not file.filename or not file.filename.lower().endswith(".pdf"):
            raise HTTPException(status_code=400, detail="Invalid file type. Please upload a PDF file.")
        if len(pdf_bytes) == 0:
            raise HTTPException(status_code=400, detail="The uploaded PDF file is empty.")

        extracted = await run_blocking(PDFService.extract_pages, pdf_bytes=pdf_bytes, pages=pages)
        base = os.path.splitext(file.filename)[0]
        output_filename = f"{base}_extracted.pdf"
        return StreamingResponse(
            io.BytesIO(extracted),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{output_filename}"',
                "Content-Length": str(len(extracted)),
            },
        )
    except HTTPException:
        raise
    except PDFProcessingError as e:
        raise HTTPException(status_code=400, detail=e.message)
    except Exception as e:
        logger.exception("[extract_pages] unexpected error")
        raise HTTPException(status_code=500, detail="An unexpected error occurred while extracting pages.")


@router.post("/page-numbering", summary="Add page numbers to PDF")
async def page_numbering(
    file: UploadFile = File(...),
    number_format: str = Form("1,2,3", description="Number format: 1,2,3 | I,II,III | i,ii,iii | A,B,C | Page 1 | 1 of 10 | PAGE-001"),
    starting_number: int = Form(1, ge=1, description="Starting page number"),
    format_template: str = Form("{n}", description="Format template with {n} and {total} placeholders"),
    position: str = Form("bottom-center", description="Position: top-left, top-center, top-right, bottom-left, bottom-center, bottom-right"),
    alignment: str = Form("center", description="Alignment: left, center, right"),
    page_range: str = Form("all", description="Page range: all, odd, even, first, 1-5, 2,4,6"),
    font_size: int = Form(12, ge=6, le=72, description="Font size (6-72)"),
    font_color: str = Form("#000000", description="Font color in hex format"),
    font_family: str = Form("Helvetica", description="Font family: Helvetica, Courier, Times-Roman"),
    prefix: str = Form("", description="Text to add before page number"),
    suffix: str = Form("", description="Text to add after page number"),
):
    """
    Add page numbers to a PDF with full customization.
    Supports multiple number formats, positions, and page range filtering.
    """
    import os
    try:
        # Read the uploaded PDF
        pdf_bytes = await file.read()

        # Validate the file is a PDF
        if not file.filename or not file.filename.lower().endswith(".pdf"):
            raise HTTPException(
                status_code=400,
                detail="Invalid file type. Please upload a PDF file.",
            )

        if len(pdf_bytes) == 0:
            raise HTTPException(
                status_code=400,
                detail="The uploaded PDF file is empty.",
            )

        # Validate font_color hex format
        if not font_color.startswith("#") or len(font_color) != 7:
            font_color = "#000000"

        logger.debug(f"[page_numbering] Processing PDF: {file.filename} "
                     f"(format={number_format}, position={position}, range={page_range}, "
                     f"font_size={font_size}, color={font_color})")

        # Apply page numbering
        numbered_pdf = await run_blocking(
            PDFService.page_numbering,
            pdf_bytes=pdf_bytes,
            number_format=number_format,
            starting_number=starting_number,
            format_template=format_template,
            position=position,
            alignment=alignment,
            page_range=page_range,
            font_size=font_size,
            font_color=font_color,
            font_family=font_family,
            prefix=prefix,
            suffix=suffix,
        )

        # Generate output filename
        base_name = os.path.splitext(file.filename)[0]
        output_filename = f"{base_name}_numbered.pdf"

        logger.debug(f"[page_numbering] Successfully added page numbers, returning "
                     f"{len(numbered_pdf)} bytes as '{output_filename}'")

        return StreamingResponse(
            io.BytesIO(numbered_pdf),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{output_filename}"',
                "Content-Length": str(len(numbered_pdf)),
            },
        )

    except HTTPException:
        raise
    except PDFProcessingError as e:
        logger.debug(f"[page_numbering] PDFProcessingError: {e.message}")
        raise HTTPException(
            status_code=500,
            detail=f"Page numbering failed: {e.message}",
        )
    except Exception as e:
        logger.debug(f"[page_numbering] Unexpected error: {type(e).__name__}: {e}")
        raise HTTPException(
            status_code=500,
            detail="An unexpected error occurred while adding page numbers.",
        )


@router.post("/organize-pdf", summary="Organize/rearrange PDF pages")
async def organize_pdf(
    file: UploadFile = File(...),
    page_order: Optional[str] = Form(None, description="JSON array of 1-based page numbers in desired order"),
    deleted_pages: Optional[str] = Form(None, description="JSON array of 1-based page numbers to delete"),
):
    """
    Reorganize PDF pages — reorder, delete, or both.

    Accepts a PDF file plus optional page_order and deleted_pages as JSON strings.
    Returns the reorganized PDF as a download.
    """
    try:
        logger.debug(f"[organize_pdf] Processing: {file.filename}, page_order={page_order}, deleted_pages={deleted_pages}")

        # Validate file
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided.")
        if not file.content_type or file.content_type != "application/pdf":
            raise HTTPException(status_code=400, detail="Only PDF files are supported.")

        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Uploaded file is empty.")
        if len(content) > 100 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File too large. Maximum size is 100MB.")

        # Parse page_order JSON
        parsed_page_order: Optional[List[int]] = None
        if page_order:
            import json
            try:
                parsed_page_order = json.loads(page_order)
                if not isinstance(parsed_page_order, list) or not all(isinstance(p, int) for p in parsed_page_order):
                    raise ValueError("page_order must be a JSON array of integers")
                logger.debug(f"[organize_pdf] Parsed page_order: {parsed_page_order}")
            except (json.JSONDecodeError, ValueError) as e:
                raise HTTPException(status_code=400, detail=f"Invalid page_order: {str(e)}")

        # Parse deleted_pages JSON
        parsed_deleted_pages: Optional[List[int]] = None
        if deleted_pages:
            import json
            try:
                parsed_deleted_pages = json.loads(deleted_pages)
                if not isinstance(parsed_deleted_pages, list) or not all(isinstance(p, int) for p in parsed_deleted_pages):
                    raise ValueError("deleted_pages must be a JSON array of integers")
                logger.debug(f"[organize_pdf] Parsed deleted_pages: {parsed_deleted_pages}")
            except (json.JSONDecodeError, ValueError) as e:
                raise HTTPException(status_code=400, detail=f"Invalid deleted_pages: {str(e)}")

        # Process via service
        output_bytes = await run_blocking(
            OrganizePDFService.organize_pdf,
            pdf_bytes=content,
            page_order=parsed_page_order,
            deleted_pages=parsed_deleted_pages,
        )

        output_filename = f"{file.filename.rsplit('.', 1)[0]}_organized.pdf"
        logger.debug(f"[organize_pdf] Successfully organized, returning {len(output_bytes)} bytes as '{output_filename}'")

        return StreamingResponse(
            io.BytesIO(output_bytes),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{output_filename}"',
                "Content-Length": str(len(output_bytes)),
            },
        )

    except ValueError as ve:
        logger.debug(f"[organize_pdf] Validation error: {ve}")
        raise HTTPException(status_code=400, detail=str(ve))
    except PDFProcessingError as e:
        logger.debug(f"[organize_pdf] Processing error: {e.message}")
        raise HTTPException(
            status_code=500,
            detail=f"PDF organization failed: {e.message}",
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.debug(f"[organize_pdf] Unexpected error: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail="An unexpected error occurred while organizing the PDF.",
        )


# ══════════════════════════════════════════════════════════════════════
# AI Tools Endpoints
# ══════════════════════════════════════════════════════════════════════

@router.post("/ai-tools", summary="Analyze PDF with AI")
async def ai_tools_analyze(file: UploadFile = File(...)):
    """
    Analyze a PDF using AI — extract text, generate summary, key points,
    title, and sentiment analysis.

    Returns JSON with full analysis results.
    """
    logger.debug("[ai-tools] Endpoint called")
    try:
        # Validate file
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        if not file.filename.lower().endswith('.pdf'):
            raise HTTPException(status_code=400, detail="Only PDF files are supported for AI analysis")

        if not validate_file_size(file, max_size=50 * 1024 * 1024):
            raise HTTPException(status_code=413, detail="File exceeds maximum size of 50MB")

        # Read file content
        pdf_bytes = await file.read()
        if not pdf_bytes:
            raise HTTPException(status_code=400, detail="Empty file uploaded")

        logger.debug(f"[ai-tools] Processing: {file.filename} ({len(pdf_bytes)} bytes)")

        # Run AI analysis
        from app.services.ai_tools_service import analyze_pdf
        async with heavy_job_slot():
            result = await run_blocking(analyze_pdf, pdf_bytes)

        logger.debug(f"[ai-tools] Analysis complete: summary={len(result['summary'])} chars, "
              f"keyPoints={len(result['keyPoints'])}, sentiment={result['sentiment']}")

        return result

    except HTTPException:
        raise
    except Exception as e:
        logger.debug(f"[ai-tools] Error: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"AI analysis failed: {str(e)}",
        )


@router.post("/ai-tools/report", summary="Generate AI analysis report (DOCX)")
async def ai_tools_report(file: UploadFile = File(...)):
    """
    Analyze a PDF with AI AND generate a downloadable DOCX report.

    Returns a .docx file with the complete AI analysis report.
    """
    logger.debug("[ai-tools/report] Endpoint called")
    try:
        # Validate file
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        if not file.filename.lower().endswith('.pdf'):
            raise HTTPException(status_code=400, detail="Only PDF files are supported for AI analysis")

        if not validate_file_size(file, max_size=50 * 1024 * 1024):
            raise HTTPException(status_code=413, detail="File exceeds maximum size of 50MB")

        # Read file content
        pdf_bytes = await file.read()
        if not pdf_bytes:
            raise HTTPException(status_code=400, detail="Empty file uploaded")

        logger.debug(f"[ai-tools/report] Processing: {file.filename} ({len(pdf_bytes)} bytes)")

        # Run AI analysis
        from app.services.ai_tools_service import analyze_pdf, generate_report
        async with heavy_job_slot():
            result = await run_blocking(analyze_pdf, pdf_bytes)

            # Generate DOCX report
            report_bytes = await run_blocking(generate_report, result, original_filename=file.filename)

        # Determine output filename
        base_name = file.filename.rsplit('.', 1)[0] if '.' in file.filename else file.filename
        download_filename = f"{base_name}_ai_report.docx"

        logger.debug(f"[ai-tools/report] Report generated: {download_filename} ({len(report_bytes)} bytes)")

        return StreamingResponse(
            io.BytesIO(report_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{download_filename}"',
                "Content-Length": str(len(report_bytes)),
            },
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.debug(f"[ai-tools/report] Error: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"AI report generation failed: {str(e)}",
        )


# ============================================================
# NEW: PDF to Excel endpoint
# ============================================================
@router.post('/pdf-to-excel', summary='Convert PDF to Excel spreadsheet')
async def pdf_to_excel(file: UploadFile = File(...)):
    """Convert PDF to Excel spreadsheet by extracting tables."""
    logger.debug('[pdf-to-excel] Endpoint called')
    if not file:
        raise HTTPException(status_code=400, detail='No file provided')

    try:
        content_type = file.content_type or ''
        filename = (file.filename or '').lower()
        is_pdf = (
            content_type == 'application/pdf'
            or filename.endswith('.pdf')
            or '.pdf' in content_type
        )

        if not is_pdf:
            raise HTTPException(
                status_code=400,
                detail=f'Invalid file type. Expected PDF but got: {content_type or filename}'
            )

        file_bytes = await file.read()
        file_size = len(file_bytes)

        if file_size > 100 * 1024 * 1024:
            raise HTTPException(status_code=413, detail='File exceeds 100MB limit')

        logger.debug(f'[pdf-to-excel] File received: {file.filename} ({file_size} bytes)')

        from app.services.pdf_to_excel_service import PdfToExcelService

        service = PdfToExcelService()
        async with heavy_job_slot():
            excel_bytes, _ = await run_blocking(
                service.convert_to_excel, file_bytes, file.filename or 'document.pdf'
            )

        output_filename = Path(file.filename or 'document').stem + '.xlsx'
        logger.debug(f'[pdf-to-excel] Conversion complete. Output size: {len(excel_bytes)} bytes')

        return create_file_response(
            file_bytes=excel_bytes,
            filename=output_filename,
            media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )

    except HTTPException:
        raise
    except ValueError as e:
        logger.debug(f'[pdf-to-excel] ValueError: {e}')
        raise HTTPException(status_code=422, detail=str(e))
    except Exception as e:
        logger.debug(f'[pdf-to-excel] Error: {type(e).__name__}: {e}')
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f'PDF-to-Excel conversion failed: {str(e)}',
        )


# ============================================================
# NEW: Excel to PDF endpoint
# ============================================================
@router.post('/excel-to-pdf', summary='Convert Excel spreadsheet to PDF')
async def excel_to_pdf(file: UploadFile = File(...)):
    """Convert Excel spreadsheet (.xlsx/.xls) to PDF."""
    logger.debug('[excel-to-pdf] Endpoint called')
    if not file:
        raise HTTPException(status_code=400, detail='No file provided')

    try:
        content_type = file.content_type or ''
        filename = (file.filename or '').lower()

        is_excel = (
            content_type in (
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                'application/vnd.ms-excel',
            )
            or (content_type == 'application/octet-stream'
                and (filename.endswith('.xlsx') or filename.endswith('.xls')))
            or filename.endswith('.xlsx')
            or filename.endswith('.xls')
        )

        if not is_excel:
            raise HTTPException(
                status_code=400,
                detail=f'Invalid file type. Expected Excel (.xlsx/.xls) but got: {content_type or filename}'
            )

        file_bytes = await file.read()
        file_size = len(file_bytes)

        if file_size > 100 * 1024 * 1024:
            raise HTTPException(status_code=413, detail='File exceeds 100MB limit')

        logger.debug(f'[excel-to-pdf] File received: {file.filename} ({file_size} bytes)')

        from app.services.excel_to_pdf_service import ExcelToPdfService

        service = ExcelToPdfService()
        async with heavy_job_slot():
            pdf_bytes, _ = await run_blocking(
                service.convert_to_pdf, file_bytes, file.filename or 'spreadsheet.xlsx'
            )

        output_filename = Path(file.filename or 'spreadsheet').stem + '.pdf'
        logger.debug(f'[excel-to-pdf] Conversion complete. Output size: {len(pdf_bytes)} bytes')

        return create_file_response(
            file_bytes=pdf_bytes,
            filename=output_filename,
            media_type='application/pdf',
        )

    except HTTPException:
        raise
    except ValueError as e:
        logger.debug(f'[excel-to-pdf] ValueError: {e}')
        raise HTTPException(status_code=422, detail=str(e))
    except Exception as e:
        logger.debug(f'[excel-to-pdf] Error: {type(e).__name__}: {e}')
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f'Excel-to-PDF conversion failed: {str(e)}',
        )


@router.post("/powerpoint-to-pdf", summary="Convert PowerPoint to PDF")
async def powerpoint_to_pdf(file: UploadFile = File(...)):
    """Convert .ppt/.pptx to PDF using LibreOffice (same engine as word-to-pdf)."""
    from app.services.pptx_service import PptxService

    if not file.filename:
        raise HTTPException(status_code=400, detail="File must have a filename.")
    name_lower = file.filename.lower()
    if not (name_lower.endswith(".ppt") or name_lower.endswith(".pptx")):
        raise HTTPException(
            status_code=415,
            detail="File must be a PowerPoint document (.ppt or .pptx).",
        )

    try:
        ppt_bytes = await read_upload_file(file)
        if not ppt_bytes:
            raise HTTPException(status_code=400, detail="Uploaded file is empty.")

        async with heavy_job_slot():
            pdf_bytes = await run_blocking(PptxService.pptx_to_pdf, ppt_bytes, file.filename)
        base = os.path.splitext(file.filename)[0]
        return create_file_response(pdf_bytes, f"{base}_converted.pdf", media_type="application/pdf")
    except HTTPException:
        raise
    except PDFProcessingError as e:
        raise HTTPException(status_code=500, detail=e.message)
    except Exception as e:
        logger.exception("[powerpoint_to_pdf] unexpected error")
        raise HTTPException(status_code=500, detail=f"PowerPoint to PDF failed: {e}")


@router.post("/pdf-to-powerpoint", summary="Convert PDF to PowerPoint")
async def pdf_to_powerpoint(file: UploadFile = File(...)):
    """Convert PDF to .pptx by rendering each page as a slide-sized image."""
    from app.services.pptx_service import PptxService

    if not file.filename or not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Invalid file type. Please upload a PDF file.")

    try:
        pdf_bytes = await file.read()
        if not pdf_bytes:
            raise HTTPException(status_code=400, detail="Uploaded PDF is empty.")

        async with heavy_job_slot():
            pptx_bytes = await run_blocking(PptxService.pdf_to_pptx, pdf_bytes)
        base = os.path.splitext(file.filename)[0]
        return create_file_response(
            pptx_bytes,
            f"{base}_converted.pptx",
            media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
        )
    except HTTPException:
        raise
    except PDFProcessingError as e:
        raise HTTPException(status_code=400, detail=e.message)
    except Exception as e:
        logger.exception("[pdf_to_powerpoint] unexpected error")
        raise HTTPException(status_code=500, detail=f"PDF to PowerPoint failed: {e}")


@router.post("/ocr-pdf", summary="Make a scanned PDF searchable with OCR")
async def ocr_pdf(
    file: UploadFile = File(...),
    language: str = Form("eng", description="OCR language code (eng or hin)"),
):
    """Add an invisible text layer to a scanned PDF using ocrmypdf."""
    from app.services.ocr_pdf_service import OcrPdfService

    if not file.filename or not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Invalid file type. Please upload a PDF file.")

    try:
        pdf_bytes = await file.read()
        if not pdf_bytes:
            raise HTTPException(status_code=400, detail="Uploaded PDF is empty.")

        async with heavy_job_slot():
            searchable = await OcrPdfService.ocr_pdf(pdf_bytes, language=language)
        base = os.path.splitext(file.filename)[0]
        return create_file_response(
            searchable,
            f"{base}_searchable.pdf",
            media_type="application/pdf",
        )
    except HTTPException:
        raise
    except PDFProcessingError as e:
        # Ghostscript missing / language unsupported is a 400, not 500.
        raise HTTPException(status_code=400, detail=e.message)
    except Exception as e:
        logger.exception("[ocr_pdf] unexpected error")
        raise HTTPException(status_code=500, detail=f"OCR failed: {e}")


@router.post("/sign-pdf", summary="Stamp a signature image onto one or more PDF positions")
async def sign_pdf(
    file: UploadFile = File(..., description="PDF to sign"),
    signature: UploadFile = File(..., description="Signature image (PNG/JPG)"),
    placements: Optional[str] = Form(
        None,
        description='JSON array of {"page_index","x","y","width","height"} objects (ratios 0..1)',
    ),
    page_index: Optional[int] = Form(None, ge=0, description="Legacy: single placement page index"),
    x: Optional[float] = Form(None, description="Legacy: single placement x ratio"),
    y: Optional[float] = Form(None, description="Legacy: single placement y ratio"),
    width: Optional[float] = Form(None, description="Legacy: single placement width ratio"),
    height: Optional[float] = Form(None, description="Legacy: single placement height ratio"),
):
    """Stamp a signature image at one or more positions across the PDF.

    Prefer the `placements` JSON array. The flat form fields are kept for
    backward compatibility with the original single-placement clients.
    """
    import json

    from app.services.sign_pdf_service import SignPdfService

    if not file.filename or not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Invalid file type. Please upload a PDF file.")

    try:
        pdf_bytes = await file.read()
        sig_bytes = await signature.read()
        if not pdf_bytes:
            raise HTTPException(status_code=400, detail="Uploaded PDF is empty.")
        if not sig_bytes:
            raise HTTPException(status_code=400, detail="Signature image is empty.")

        parsed: list[tuple[int, float, float, float, float]] = []
        if placements:
            try:
                raw = json.loads(placements)
            except json.JSONDecodeError as e:
                raise HTTPException(status_code=400, detail=f"placements is not valid JSON: {e}")
            if not isinstance(raw, list) or not raw:
                raise HTTPException(status_code=400, detail="placements must be a non-empty array.")
            for i, item in enumerate(raw):
                if not isinstance(item, dict):
                    raise HTTPException(status_code=400, detail=f"placements[{i}] must be an object.")
                try:
                    parsed.append((
                        int(item["page_index"]),
                        float(item["x"]),
                        float(item["y"]),
                        float(item["width"]),
                        float(item["height"]),
                    ))
                except (KeyError, TypeError, ValueError) as e:
                    raise HTTPException(
                        status_code=400,
                        detail=f"placements[{i}] missing/invalid field: {e}",
                    )
        else:
            if None in (page_index, x, y, width, height):
                raise HTTPException(
                    status_code=400,
                    detail="Provide either `placements` JSON or all of page_index/x/y/width/height.",
                )
            parsed.append((int(page_index), float(x), float(y), float(width), float(height)))

        signed = await run_blocking(
            SignPdfService.sign_pdf,
            pdf_bytes=pdf_bytes,
            signature_png_bytes=sig_bytes,
            placements=parsed,
        )
        base = os.path.splitext(file.filename)[0]
        return create_file_response(signed, f"{base}_signed.pdf", media_type="application/pdf")
    except HTTPException:
        raise
    except PDFProcessingError as e:
        raise HTTPException(status_code=400, detail=e.message)
    except Exception as e:
        logger.exception("[sign_pdf] unexpected error")
        raise HTTPException(status_code=500, detail=f"Signing failed: {e}")