"""
AI Tools Service — PDF text extraction, summarization, key points, title generation,
sentiment analysis, and report generation.

Uses:
- pymupdf (fitz) for high-quality PDF text extraction
- HuggingFace transformers for NLP (bart-large-cnn, sentiment-analysis)
- python-docx for report generation
"""

import io
import re
import logging
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# ── Lazy-loaded heavy imports ────────────────────────────────────────
import fitz  # pymupdf — always available (lightweight PDF parsing)
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Cloud LLM (Groq) — primary AI engine when configured. All public functions
# return Optional and we fall back to local HuggingFace pipelines on None.
from app.services import cloud_ai_service

logger = logging.getLogger("ai_tools")

# HuggingFace pipelines will be lazily loaded per-call to avoid blocking
# the event loop on first import (model download / warm-up can take time).
# For production, pre-load in a startup event.

_SUMMARIZER = None
_SENTIMENT_PIPELINE = None

MODEL_SUMMARY = "facebook/bart-large-cnn"  # good summarization model
CHUNK_MAX_CHARS = 1024  # Safe token limit for bart-large-cnn


# ── Lazy initializers ────────────────────────────────────────────────

def _get_summarizer():
    """Lazily load the summarization pipeline."""
    global _SUMMARIZER
    if _SUMMARIZER is None:
        from transformers import pipeline
        logger.info("Loading summarization model: %s", MODEL_SUMMARY)
        _SUMMARIZER = pipeline("summarization", model=MODEL_SUMMARY)
    return _SUMMARIZER


def _get_sentiment_pipeline():
    """Lazily load the sentiment-analysis pipeline."""
    global _SENTIMENT_PIPELINE
    if _SENTIMENT_PIPELINE is None:
        from transformers import pipeline
        logger.info("Loading sentiment-analysis model")
        _SENTIMENT_PIPELINE = pipeline("sentiment-analysis")
    return _SENTIMENT_PIPELINE


# ── Text extraction ──────────────────────────────────────────────────

def extract_pdf_text(pdf_bytes: bytes) -> Tuple[str, int]:
    """Extract full text from a PDF using pymupdf.

    Returns:
        (extracted_text, page_count)
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page_count = doc.page_count
    full_text_parts: List[str] = []

    for page_num in range(page_count):
        page = doc[page_num]
        page_text = page.get_text("text")  # plain text with layout
        full_text_parts.append(page_text)

    doc.close()
    full_text = "\n\n".join(full_text_parts)
    logger.info("Extracted %d chars of text from %d pages", len(full_text), page_count)
    return full_text, page_count


def _compute_reading_time(word_count: int) -> str:
    """Estimate reading time in minutes (average 200 wpm)."""
    minutes = max(1, round(word_count / 200))
    if minutes < 1:
        return "< 1 minute"
    elif minutes == 1:
        return "1 minute"
    else:
        return f"{minutes} minutes"


# ── Text chunking for models with token limits ───────────────────────

def _chunk_text(text: str, max_chars: int = CHUNK_MAX_CHARS) -> List[str]:
    """Split text into roughly equal chunks, respecting sentence boundaries."""
    if len(text) <= max_chars:
        return [text]

    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks: List[str] = []
    current_chunk: List[str] = []
    current_len = 0

    for sentence in sentences:
        if current_len + len(sentence) > max_chars and current_chunk:
            chunks.append(" ".join(current_chunk))
            current_chunk = [sentence]
            current_len = len(sentence)
        else:
            current_chunk.append(sentence)
            current_len += len(sentence)

    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return chunks if chunks else [text]


# ── AI Analysis ──────────────────────────────────────────────────────

def _generate_summary(text: str, max_length: int = 150, min_length: int = 40) -> str:
    """Generate a concise summary using cloud LLM (Groq) with HuggingFace fallback."""
    # Try cloud LLM first — higher quality, faster
    cloud_result = cloud_ai_service.summarize(text, max_words=max_length // 2)
    if cloud_result:
        logger.info("Summary: cloud LLM (Groq)")
        return cloud_result

    # Fallback: HuggingFace BART (slower, lower quality)
    logger.info("Summary: HuggingFace fallback")
    # Truncate if text is too long
    input_text = text[:4000] if len(text) > 4000 else text
    if len(input_text.strip()) < 50:
        return input_text.strip()

    try:
        summarizer = _get_summarizer()
        result = summarizer(
            input_text,
            max_length=min(max_length, 200),
            min_length=min(min_length, 30),
            do_sample=False,
        )
        if result and len(result) > 0:
            return result[0]["summary_text"].strip()
    except Exception as e:
        logger.warning("Summarization failed: %s", e)

    # Last-resort fallback: return first few sentences
    sentences = re.split(r'(?<=[.!?])\s+', input_text)
    return " ".join(sentences[:3])


def _extract_key_points(text: str, num_points: int = 5) -> List[str]:
    """Extract key points using cloud LLM with HuggingFace + heuristic fallback."""
    if len(text.strip()) < 50:
        return ["Text too short for key point extraction."]

    # Try cloud LLM first — gives genuinely distinct insights, not duplicates
    cloud_points = cloud_ai_service.extract_key_points(text, num_points=num_points)
    if cloud_points:
        logger.info("KeyPoints: cloud LLM (Groq), %d points", len(cloud_points))
        return cloud_points

    # Fallback: HuggingFace per-chunk summarization
    logger.info("KeyPoints: HuggingFace fallback")
    chunks = _chunk_text(text, max_chars=1200)
    all_points: List[str] = []

    # Summarize each chunk to get a mini-insight
    for chunk in chunks[:5]:  # limit to avoid excessive API calls
        if len(chunk.strip()) < 20:
            continue
        try:
            mini_summary = _generate_summary(chunk, max_length=80, min_length=20)
            if mini_summary and mini_summary not in all_points:
                all_points.append(mini_summary)
        except Exception:
            pass

    if len(all_points) < num_points:
        # Last-resort: extract longest sentences as points
        sentences = re.split(r'(?<=[.!?])\s+', text)
        long_sentences = sorted(
            [s.strip() for s in sentences if len(s.strip()) > 30],
            key=len, reverse=True
        )
        for s in long_sentences:
            if s not in all_points and len(all_points) < num_points:
                all_points.append(s)

    return all_points[:num_points]


def _generate_title(text: str) -> str:
    """Generate a smart title using cloud LLM with HuggingFace fallback."""
    if len(text.strip()) < 30:
        return "Untitled Document"

    # Try cloud LLM first — produces clean, properly-formatted titles
    cloud_title = cloud_ai_service.generate_title(text)
    if cloud_title:
        logger.info("Title: cloud LLM (Groq)")
        return cloud_title

    # Fallback: HuggingFace BART
    logger.info("Title: HuggingFace fallback")
    try:
        summarizer = _get_summarizer()
        # Use a very short summary as title
        result = summarizer(
            text[:3000],
            max_length=30,
            min_length=5,
            do_sample=False,
        )
        if result and len(result) > 0:
            title = result[0]["summary_text"].strip()
            # Clean up: capitalize words, remove trailing periods
            title = title.rstrip(".")
            title = " ".join(
                w.capitalize() if len(w) > 2 else w for w in title.split()
            )
            return title
    except Exception as e:
        logger.warning("Title generation failed: %s", e)

    # Last-resort: use first sentence
    sentences = re.split(r'(?<=[.!?])\s+', text)
    first = sentences[0].strip().rstrip(".")
    if len(first) > 60:
        first = first[:60] + "..."
    return first if first else "Document Analysis"


def _analyze_sentiment(text: str) -> Tuple[str, float]:
    """Analyze sentiment using cloud LLM with HuggingFace fallback."""
    if len(text.strip()) < 20:
        return "neutral", 50.0

    # Try cloud LLM first — better nuance than binary distilbert
    cloud_result = cloud_ai_service.analyze_sentiment(text)
    if cloud_result:
        logger.info("Sentiment: cloud LLM (Groq)")
        return cloud_result

    # Fallback: HuggingFace distilbert (positive/negative only, no neutral)
    logger.info("Sentiment: HuggingFace fallback")
    sample = text[:1000] if len(text) > 1000 else text
    try:
        pipeline = _get_sentiment_pipeline()
        result = pipeline(sample)
        if result and len(result) > 0:
            label = result[0]["label"].lower()
            score = round(result[0]["score"] * 100, 1)
            if label in ("positive", "pos", "5 stars", "4 stars"):
                return "positive", score
            elif label in ("negative", "neg", "1 star", "2 stars"):
                return "negative", score
            else:
                return "neutral", score
    except Exception as e:
        logger.warning("Sentiment analysis failed: %s", e)

    return "neutral", 50.0


# ── Word count ───────────────────────────────────────────────────────

def _count_words(text: str) -> int:
    """Count words in text."""
    return len(re.findall(r'\b\w+\b', text))


# ── Full analysis pipeline ───────────────────────────────────────────

def analyze_pdf(pdf_bytes: bytes) -> Dict:
    """Run the complete AI analysis pipeline on a PDF.

    Returns a dict suitable for JSON serialization:
    {
        "summary": str,
        "keyPoints": [str, ...],
        "title": str,
        "wordCount": int,
        "pageCount": int,
        "readingTime": str,
        "sentiment": "positive" | "neutral" | "negative",
        "confidence": float (0-100),
    }
    """
    logger.info("=" * 50)
    logger.info("Starting AI analysis pipeline")

    # 1. Extract text
    logger.info("Extracting PDF text...")
    text, page_count = extract_pdf_text(pdf_bytes)

    if not text.strip():
        return {
            "summary": "No extractable text found in this PDF. The document may be scanned/image-based.",
            "keyPoints": ["No text content to analyze."],
            "title": "Untitled Document",
            "wordCount": 0,
            "pageCount": page_count,
            "readingTime": "< 1 minute",
            "sentiment": "neutral",
            "confidence": 0,
        }

    word_count = _count_words(text)
    logger.info("Word count: %d", word_count)

    # 2. Summarize
    logger.info("Generating summary...")
    summary = _generate_summary(text)

    # 3. Key points
    logger.info("Extracting key points...")
    key_points = _extract_key_points(text)

    # 4. Title
    logger.info("Generating title...")
    title = _generate_title(text)

    # 5. Sentiment
    logger.info("Analyzing sentiment...")
    sentiment_label, confidence = _analyze_sentiment(text)

    # 6. Reading time
    reading_time = _compute_reading_time(word_count)

    result = {
        "summary": summary,
        "keyPoints": key_points,
        "title": title,
        "wordCount": word_count,
        "pageCount": page_count,
        "readingTime": reading_time,
        "sentiment": sentiment_label,
        "confidence": confidence,
    }

    logger.info("AI analysis completed successfully")
    logger.info("=" * 50)
    return result


# ── Report generation (DOCX) ─────────────────────────────────────────

def generate_report(
    analysis_result: Dict,
    original_filename: str = "document.pdf",
) -> bytes:
    """Generate a formatted DOCX report from AI analysis results.

    Returns the DOCX file as bytes.
    """
    doc = DocxDocument()

    # -- Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # -- Title
    title_para = doc.add_heading("AI Analysis Report", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Source: {original_filename}")
    doc.add_paragraph(f"Generated by WeLovePDF AI Tools")
    doc.add_paragraph("─" * 50)

    # -- Document Statistics
    doc.add_heading("Document Statistics", level=1)
    stats_table = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
    stats_data = [
        ("Pages", str(analysis_result.get("pageCount", "N/A"))),
        ("Words", f"{analysis_result.get('wordCount', 0):,}"),
        ("Reading Time", analysis_result.get("readingTime", "N/A")),
        ("Sentiment", analysis_result.get("sentiment", "neutral").title()),
        ("AI Confidence", f"{analysis_result.get('confidence', 0):.1f}%"),
    ]
    for i, (label, value) in enumerate(stats_data):
        stats_table.cell(i, 0).text = label
        stats_table.cell(i, 1).text = value

    doc.add_paragraph()

    # -- Generated Title
    doc.add_heading("Generated Title", level=1)
    doc.add_paragraph(analysis_result.get("title", "Untitled"))

    doc.add_paragraph()

    # -- Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(analysis_result.get("summary", "No summary available."))

    doc.add_paragraph()

    # -- Key Points
    doc.add_heading("Key Points", level=1)
    key_points = analysis_result.get("keyPoints", [])
    if key_points:
        for i, point in enumerate(key_points, 1):
            doc.add_paragraph(f"{i}. {point}", style="List Number")
    else:
        doc.add_paragraph("No key points extracted.")

    doc.add_paragraph()

    # -- Footer
    doc.add_paragraph("─" * 50)
    footer = doc.add_paragraph("Report generated by WeLovePDF AI Tools.")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()